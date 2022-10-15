# imports
from docx import Document
from docx.shared import Inches,Pt,Cm
from docx2pdf import convert
import os
def createpdfs_grade9(parameter_name, parameter_data, parameter_filename):
    # initialise
    document = Document()

    # variables
    name = parameter_name
    filename = parameter_filename
    # changed to 5 subjects for now
    data1 = list(parameter_data.keys())
    subject1, subject2, subject3, subject4, subject5 = data1[3], data1[4], data1[5], data1[6], data1[7]
    data2 = list(parameter_data.values())
    percentage1, percentage2, percentage3, percentage4, percentage5 = data2[3], data2[4], data2[5], data2[6], data2[7]
    date = 'September 30, 2022'
    session = "2021-22"
    teacher_name = "Sana Kapur"
    grade = "IX"

    # document build
    p = document.add_paragraph("")
    document.add_picture('resources/header.png', width=Inches(6))

    p = document.add_paragraph("")
    run = p.add_run(date)
    p.alignment = 2
    run.bold = True

    p = document.add_paragraph("")
    run = p.add_run('TRANSCRIPT FOR '+name+'')
    run.bold = True
    run.underline = True
    p.alignment = 1

    p = document.add_paragraph('Sanskriti School is a private high school in New Delhi, India. We have students in the age group of 3 years to 17 years. Our total strength is approximately 2500 students, of these there are approximately 250 students in the graduating class.' )
    p = document.add_paragraph('In the academic system that we follow, we do not rank the students. The students usually take a Unit Test once a week and there are term end exams. The academic performance is based on the online assessment.' )

    p = document.add_paragraph("")
    run = p.add_run('Following is the final percentage for grade '+grade+' (Academic Session: '+session+'): ')
    run.bold = True

    records = (
        (subject1, round(float(percentage1)*100/190)),
        (subject2, round(float(percentage2)*100/190)),
        (subject3, round(float(percentage3)*100/190)),
        (subject4, round(float(percentage4)*100/190)),
        (subject5, round(float(percentage5)*100/190))

        )

    table = document.add_table(rows=1, cols=2)
    table.style= "Table Grid"
    hdr_cells = table.rows[0].cells

    hdr_cells[0].text = 'Subjects'
    hdr_cells[1].text = "Grade " + grade
    hdr_cells[1].paragraphs[0].alignment = 1
    hdr_cells[1].paragraphs[0].runs[0].bold = True
    hdr_cells[0].paragraphs[0].runs[0].bold = True


    for subj, percentage in records:
        row_cells = table.add_row().cells

        row_cells[0].text = str(subj)
        row_cells[1].text = str(percentage)+"%"
        row_cells[1].paragraphs[0].alignment = 1

    p = document.add_paragraph("")
    # table.cell(0,0).vertical_alignment = 1

    paragraph = hdr_cells[0].paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(12)
    paragraph = hdr_cells[1].paragraphs[0]
    run = paragraph.runs
    font = run[0].font
    font.size = Pt(12)
    for row in table.rows:
        row.height = Cm(0.7)
    document.add_picture('resources/sign.png', width=Inches(1))

    p = document.add_paragraph("")
    run = p.add_run(''''''+teacher_name+'''
    Senior School Counselor''')
    run.bold = True

    # save as docx
    document.save('results/demo.docx')
    # convert to pdf
    convert("results/demo.docx", "results/"+filename+".pdf")
    # remove docx
    os.remove('results/demo.docx')
